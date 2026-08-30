import io
import re
from typing import Optional
from pypdf import PdfReader
import docx

class DocumentParserService:
    """Service to parse resumes in PDF, DOCX, and TXT formats."""

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text from unnecessary whitespace and control characters."""
        if not text:
            return ""
        # Normalize multiple spaces, tabs, and carriage returns
        text = re.sub(r'\r\n|\r', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        # Normalize excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    @classmethod
    def parse_pdf(cls, file_bytes: bytes) -> str:
        """Extract text from PDF file bytes."""
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_parts.append(extracted)
            full_text = "\n".join(text_parts)
            return cls.clean_text(full_text)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @classmethod
    def parse_docx(cls, file_bytes: bytes) -> str:
        """Extract text from DOCX file bytes."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            full_text = "\n".join(text_parts)
            return cls.clean_text(full_text)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    @classmethod
    def parse_txt(cls, file_bytes: bytes) -> str:
        """Extract text from TXT file bytes with encoding detection."""
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_bytes.decode(encoding)
                return cls.clean_text(text)
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode text file with standard encodings.")

    @classmethod
    def parse_file(cls, filename: str, file_bytes: bytes) -> str:
        """Parse file according to its extension."""
        ext = filename.lower().split('.')[-1]
        if ext == 'pdf':
            return cls.parse_pdf(file_bytes)
        elif ext in ['docx', 'doc']:
            return cls.parse_docx(file_bytes)
        elif ext in ['txt', 'md']:
            return cls.parse_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file format '.{ext}'. Please upload a PDF, DOCX, or TXT file.")
